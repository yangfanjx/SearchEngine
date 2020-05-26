#!/usr/bin/python
# -*- coding: utf-8 -*-

# import docx
from docx import Document  # 导入库
import re


class Word_Table(object):

    # 广东
    @staticmethod
    def GuangDong_table(file_path):
        document = Document(file_path)  # 读入文件
        tables = document.tables  # 获取文件中的表格集
        table = tables[0]  # 获取文件中的第一个表格

        day_list = ["", ""]
        for _index, j in enumerate(table.row_cells(0)):
            if _index in (0, 1):
                continue
            day_list.append(j.text)
        print(day_list)
        time_list = ["", ""]
        for _index, j in enumerate(table.row_cells(1)):
            if _index in (0, 1):
                continue
            time_list.append(j.text)
        print(time_list)
        # title_list = ["", ""]
        # for _index, j in enumerate(table.row_cells(2)):
        #     if _index in (0, 1):
        #         continue
        #         title_list.append(j.text)
        structured_dict = {}
        for i in range(3, len(table.rows)):  # 从表格第二行开始循环读取表格数据
            m_dict = {"special_code": "",
                      "special_name": "",
                      "education": "",
                      "college": "",
                      "date": "",
                      "time": "",
                      "class_code": "",
                      "class_name": ""}

            for _index, j in enumerate(table.row_cells(i)):
                if _index == 0:
                    try:
                        _split = j.text.split()
                        if len(_split) == 3:
                            m_dict["special_code"], m_dict["special_name"], _education = _split
                            m_dict["education"] = _education[1:-1]
                        elif len(_split) == 2 and "（" in _split[1]:
                            m_dict["special_code"] = _split[0]
                            m_dict["special_name"], _education = _split[1].split("（")
                            m_dict["education"] = _education[:-1]
                    except:
                        print("error:" + str(_split))
                        continue

                elif _index == 1:
                    m_dict["college"] = j.text


                elif _index % 2 == 0:
                    m_dict["class_code"] = j.text
                else:
                    m_dict["class_name"] = j.text
                    m_dict["date"] = day_list[_index]
                    m_dict["time"] = time_list[_index]

            structured_dict[m_dict["special_code"]] = m_dict

        print(structured_dict)


    # 湖北
    @staticmethod
    def HuBei_table(file_path):
        document = Document(file_path)  # 读入文件
        tables = document.tables  # 获取文件中的表格集
        for table in tables:
            # table = tables[0]  # 获取文件中的第一个表格

            day_list = []
            for _index, j in enumerate(table.row_cells(0)):
                if _index in (0, 1, 2):
                    continue
                day_list.append(j.text)
            print(day_list)
            time_list = []
            for _index, j in enumerate(table.row_cells(1)):
                if _index > 2 and (_index - 2) % 2 == 0:
                    time_list.append(j.text)
            print(time_list)

            structured_dict = {}
            current_special = ""
            for i in range(2, len(table.rows)):  # 从表格第二行开始循环读取表格数据
                _special_code = _special_name = _class_code = ""
                for _index, j in enumerate(table.row_cells(i)):
                    if _index == 0:
                        _special_code = j.text if j.text else "公共基础课"

                        if not current_special:
                            msg_list = []
                            current_special = _special_code
                        elif current_special != _special_code:
                            # 此行专业变动
                            get_list = structured_dict.get(current_special, [])
                            get_list += msg_list
                            structured_dict[current_special] = get_list
                            current_special = _special_code
                        continue
                    elif _index == 1:
                        _special_name = j.text if j.text else "公共基础课"
                        continue
                    elif _index == 2:
                        continue
                    elif j.text:
                        if (_index - 3) % 2 == 0:
                            _class_code = j.text
                        else:
                            m_dict = {}
                            m_dict["special_code"] = _special_code
                            m_dict["special_name"] = _special_name
                            m_dict["class_code"] = _class_code
                            m_dict["class_name"] = j.text
                            m_dict["date"] = day_list[_index - 3]
                            m_dict["time"] = time_list[((_index - 3) // 2) % 2]
                            msg_list.append(m_dict)
                    else:
                        continue

                get_list = structured_dict.get(current_special, [])
                get_list += msg_list
                structured_dict[current_special] = get_list

            print(structured_dict)


    # 吉林
    @staticmethod
    def JiLin_table(file_path):

        def get_dict(_msg):
            _msg = _msg.replace(" ", "").replace(u'\n', "")
            return_list = []
            for _item in zip(re.findall(r"\d{5}", _msg), re.split(r"\d{5}", _msg)):
                return_list.append([_item[0], _item[1]])
            return return_list


        document = Document(file_path)  # 读入文件
        tables = document.tables  # 获取文件中的表格集
        for table in tables:
            # table = tables[0]  # 获取文件中的第一个表格
            day_list = []
            for _index, j in enumerate(table.row_cells(0)):
                if _index == 0:
                    continue
                if j.text:
                    day_list.append(j.text.replace(" ", "").replace(u'\n', ""))

            print(day_list)
            time_list = []
            for _index, j in enumerate(table.row_cells(1)):
                if _index == 0:
                    continue
                if j.text:
                    time_list.append(j.text.replace(" ", "").replace(u'\n', ""))
            print(time_list)

            structured_dict = {}
            _special_code = ""
            re_msg = r"\d{6}"
            for i in range(2, len(table.rows)):  # 从表格第二行开始循环读取表格数据
                class_obj_list = []
                for _index, j in enumerate(table.row_cells(i)):
                    use_msg = j.text.replace(" ", "").replace(u'\n', "")
                    if _index == 0:
                        _code_list = re.findall(re_msg, use_msg)
                        if _code_list:
                            _special_code = _code_list[0]
                            _special_name = re.sub(re_msg, "", use_msg)
                        else:
                            if "(" in use_msg or "（" in use_msg:
                                break
                            _special_code = _special_name = use_msg

                    elif use_msg:
                        for _i in get_dict(use_msg):
                            class_i = {}
                            class_i["class_code"] = _i[0]
                            class_i["class_name"] = _i[1]
                            class_i["date"] = day_list[_index - 1]
                            class_i["time"] = time_list[_index - 1]
                            class_i["special_code"] = _special_code
                            class_i["special_name"] = _special_name
                            class_obj_list.append(class_i)
                    else:
                        continue

                structured_list = structured_dict.get(_special_code, [])
                structured_list += class_obj_list
                structured_dict[_special_code] = structured_list

            print(structured_dict)


if __name__ == "__main__":
    Word_Table.JiLin_table("./download_file/1.2020年4月吉林省高等教育自学考试开考专业课程安排表.docx")
